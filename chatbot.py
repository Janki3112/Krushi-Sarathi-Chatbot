import os
import logging
import warnings
import asyncio
import time
import hashlib
import pickle
import json
from typing import List, Dict, Optional, Tuple, Any, Union
from functools import lru_cache, wraps
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading
from pathlib import Path
from datetime import datetime, timedelta
import re
from collections import defaultdict
import uuid
from concurrent.futures import ThreadPoolExecutor
from docx import Document as DocxDocument
from PyPDF2 import PdfReader

# Redis imports
import concurrent
import redis
from redis.exceptions import RedisError, ConnectionError as RedisConnectionError
import redis.sentinel
from googletrans import Translator

# LangChain imports
from langchain_groq import ChatGroq
from langchain_community.tools.tavily_search import TavilySearchResults
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.runnables.history import RunnableWithMessageHistory
from langchain_core.chat_history import BaseChatMessageHistory
from langchain_redis import RedisChatMessageHistory
from langchain_core.messages.ai import AIMessage
from langchain_core.messages.human import HumanMessage
from langchain_core.messages import messages_from_dict, messages_to_dict

# Document processing imports
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import (
    PyPDFLoader, TextLoader, UnstructuredWordDocumentLoader,
    CSVLoader, UnstructuredExcelLoader, UnstructuredMarkdownLoader
)
from langchain_community.vectorstores import Chroma
from langchain_core.runnables import RunnablePassthrough
from langchain_core.documents import Document
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity

# Website scraping imports
import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse, urlencode
from urllib.robotparser import RobotFileParser
import html2text
from dotenv import load_dotenv
load_dotenv()

# Tavily search imports
try:
    from tavily import TavilyClient
    WEB_SEARCH_AVAILABLE = True
except ImportError:
    print("⚠️ tavily-python not installed. Web search disabled.")
    print("Install with: pip install tavily-python")
    WEB_SEARCH_AVAILABLE = False

# Streamlit imports
try:
    import streamlit as st
    STREAMLIT_AVAILABLE = True
except ImportError:
    print("⚠️ streamlit not installed. UI disabled.")
    print("Install with: pip install streamlit")
    STREAMLIT_AVAILABLE = False

# Environment setup
os.environ["ANONYMIZED_TELEMETRY"] = "False"
warnings.filterwarnings("ignore")
logging.getLogger("transformers").setLevel(logging.ERROR)

# API Keys
GROQ_API_KEY = os.environ.get("GROQ_API_KEY")
TAVILY_API_KEY = os.environ.get("TAVILY_API_KEY")

# Redis Configuration
REDIS_HOST = os.environ.get("REDIS_HOST", "localhost")
REDIS_PORT = int(os.environ.get("REDIS_PORT", "6379"))
REDIS_PASSWORD = None #os.environ.get("REDIS_PASSWORD")
REDIS = int(os.environ.get("REDIS", "0"))

from dataclasses import dataclass
@dataclass

class SearchResult:
    """Data class for search results"""
    content: str
    metadata: Dict[str, Any]
    relevance_score: float = 0.0

DEFAULT_DOCS_FOLDER = os.path.join(os.path.dirname(__file__), "doc")

class OptimizedRedisManager:
    """Fixed Redis connection manager with proper error handling"""

    def __init__(
        self,
        host: str = REDIS_HOST,
        port: int = REDIS_PORT,
        password: Optional[str] = REDIS_PASSWORD,
        db: int = REDIS,
    ):
        self.host = host
        self.port = port
        self.password = password
        self.db = db
        self.redis_client = None
        self._initialize_connection()

    def _initialize_connection(self):
        try:
            redis_url = os.getenv("REDIS_URL")
            if redis_url:
                self.redis_client = redis.from_url(redis_url, decode_responses=True)
            else:
                self.redis_client = redis.Redis(
                    host=self.host,
                    port=self.port,
                    password=self.password,
                    decode_responses=True
                )
            self.redis_client.ping()
            print("✅ Redis connected")
        except Exception as e:
            print(f"❌ Redis error: {e}")
            self.redis_client = None


    def get_client(self) -> Optional[redis.Redis]:
        """Get Redis client with health check"""
        if not self.redis_client:
            return None
            
        try:
            self.redis_client.ping()
            return self.redis_client
        except Exception as e:
            print(f"🔄 Redis connection lost, attempting reconnection: {e}")
            self._initialize_connection()
            return self.redis_client

    def is_available(self) -> bool:
        """Check if Redis is available"""
        return self.redis_client is not None

class UnifiedCache:
    """Fixed unified cache with proper fallback mechanisms"""
    
    def __init__(self, redis_manager: OptimizedRedisManager, default_ttl: int = 3600):
        self.redis_manager = redis_manager
        self.default_ttl = default_ttl
        
        # Local fallback cache
        self.local_cache = {}
        self.local_cache_ttl = 300  # 5 minutes
        self.local_cache_timestamps = {}
        self.cache_hits = 0
        self.cache_misses = 0
        self.local_cache_max = 1000
        
        # Performance tracking
        self.redis_hits = 0
        self.redis_misses = 0
    
    def _get_key(self, key: str, namespace: str = "unified") -> str:
        """Generate cache key with proper namespace"""
        return f"rag_unified:{namespace}:{hashlib.md5(key.encode()).hexdigest()[:16]}"
    
    def _check_local_cache(self, cache_key: str) -> Tuple[bool, Any]:
        """Check local in-memory cache first"""
        if cache_key in self.local_cache:
            timestamp = self.local_cache_timestamps.get(cache_key, 0)
            if time.time() - timestamp < self.local_cache_ttl:
                self.cache_hits += 1
                return True, self.local_cache[cache_key]
        
        self.cache_misses += 1
        return False, None
    
    def _update_local_cache(self, cache_key: str, value: Any):
        """Update local cache with LRU eviction"""
        # LRU eviction
        if len(self.local_cache) >= self.local_cache_max:
            oldest_keys = sorted(
                self.local_cache_timestamps.keys(), 
                key=lambda k: self.local_cache_timestamps[k]
            )[:100]  # Remove oldest 100 items
            
            for key in oldest_keys:
                self.local_cache.pop(key, None)
                self.local_cache_timestamps.pop(key, None)
        
        self.local_cache[cache_key] = value
        self.local_cache_timestamps[cache_key] = time.time()
    
    def get(self, key: str, namespace: str = "unified") -> Optional[Any]:
        """Enhanced get with proper fallback"""
        cache_key = self._get_key(key, namespace)
        
        # Check local cache first
        found, value = self._check_local_cache(cache_key)
        if found:
            return value
        
        # Try Redis if available
        if self.redis_manager.is_available():
            client = self.redis_manager.get_client()
            if client:
                try:
                    cached_data = client.get(cache_key)
                    if cached_data:
                        value = json.loads(cached_data)
                        self._update_local_cache(cache_key, value)
                        self.redis_hits += 1
                        return value
                    else:
                        self.redis_misses += 1
                except Exception as e:
                    print(f"❌ Redis get error: {e}")
        
        return None
    
    def set(self, key: str, value: Any, ttl: Optional[int] = None, namespace: str = "unified") -> bool:
        """Enhanced set with proper fallback"""
        cache_key = self._get_key(key, namespace)
        success = False
        
        # Always update local cache
        self._update_local_cache(cache_key, value)
        
        # Try Redis if available
        if self.redis_manager.is_available():
            client = self.redis_manager.get_client()
            if client:
                try:
                    serialized = json.dumps(value, ensure_ascii=False, default=str)
                    ttl = ttl or self.default_ttl
                    success = client.setex(cache_key, ttl, serialized)
                except Exception as e:
                    print(f"❌ Redis set error: {e}")
        
        return success
    
    def delete(self, key: str, namespace: str = "unified") -> bool:
        """Enhanced delete with proper cleanup"""
        cache_key = self._get_key(key, namespace)
        
        # Remove from local cache
        self.local_cache.pop(cache_key, None)
        self.local_cache_timestamps.pop(cache_key, None)
        
        # Remove from Redis if available
        if self.redis_manager.is_available():
            client = self.redis_manager.get_client()
            if client:
                try:
                    return client.delete(cache_key) > 0
                except Exception:
                    pass
        
        return True  # Local removal successful
        
    def get_stats(self) -> Dict[str, Any]:
        """Enhanced cache statistics"""
        total_requests = self.cache_hits + self.cache_misses
        hit_rate = (self.cache_hits / total_requests * 100) if total_requests > 0 else 0
        
        redis_total = self.redis_hits + self.redis_misses
        redis_hit_rate = (self.redis_hits / redis_total * 100) if redis_total > 0 else 0
        
        return {
            "redis_available": self.redis_manager.is_available(),
            "local_cache_hits": self.cache_hits,
            "local_cache_misses": self.cache_misses,
            "local_hit_rate": f"{hit_rate:.1f}%",
            "redis_hits": self.redis_hits,
            "redis_misses": self.redis_misses,
            "redis_hit_rate": f"{redis_hit_rate:.1f}%",
            "local_cache_size": len(self.local_cache)
        }

class UserSearchTracker:
    """Track user searches and preferences in Redis"""
    
    def __init__(self, redis_manager: OptimizedRedisManager):
        self.redis_manager = redis_manager
        self.default_ttl = 2592000  # 30 days for user data
    
    def _get_user_searches_key(self, user_id: str) -> str:
        """Generate user searches key"""
        return f"user_searches:{user_id}"
    
    def _get_user_preferences_key(self, user_id: str) -> str:
        """Generate user preferences key"""
        return f"user_preferences:{user_id}"
    
    def _get_user_stats_key(self, user_id: str) -> str:
        """Generate user stats key"""
        return f"user_stats:{user_id}"
    
    def log_search(self, user_id: str, query: str, results_count: int, response_time: float) -> bool:
        """Log user search activity"""
        if not self.redis_manager.is_available():
            return False
        
        client = self.redis_manager.get_client()
        if not client:
            return False
        
        try:
            search_entry = {
                "query": query,
                "timestamp": datetime.now().isoformat(),
                "results_count": results_count,
                "response_time": response_time
            }
            
            # Add to user's search history (keep last 100 searches)
            searches_key = self._get_user_searches_key(user_id)
            client.lpush(searches_key, json.dumps(search_entry))
            client.ltrim(searches_key, 0, 99)  # Keep only last 100
            client.expire(searches_key, self.default_ttl)
            
            # Update user stats
            self._update_user_stats(user_id, query, results_count, response_time)
            
            return True
            
        except Exception as e:
            print(f"❌ Failed to log search: {e}")
            return False
    
    def _update_user_stats(self, user_id: str, query: str, results_count: int, response_time: float):
        """Update user statistics"""
        try:
            client = self.redis_manager.get_client()
            if not client:
                return
            
            stats_key = self._get_user_stats_key(user_id)
            
            # Get current stats
            current_stats = client.hgetall(stats_key)
            
            total_searches = int(current_stats.get('total_searches', 0)) + 1
            total_response_time = float(current_stats.get('total_response_time', 0)) + response_time
            avg_response_time = total_response_time / total_searches
            
            # Update stats
            stats_update = {
                'total_searches': total_searches,
                'total_response_time': total_response_time,
                'avg_response_time': round(avg_response_time, 2),
                'last_search': datetime.now().isoformat(),
                'last_query': query,
                'last_results_count': results_count
            }
            
            client.hset(stats_key, mapping=stats_update)
            client.expire(stats_key, self.default_ttl)
            
        except Exception as e:
            print(f"❌ Failed to update user stats: {e}")
    
    def get_user_search_history(self, user_id: str, limit: int = 20) -> List[Dict]:
        """Get user's recent search history"""
        if not self.redis_manager.is_available():
            return []
        
        client = self.redis_manager.get_client()
        if not client:
            return []
        
        try:
            searches_key = self._get_user_searches_key(user_id)
            search_entries = client.lrange(searches_key, 0, limit - 1)
            
            return [json.loads(entry) for entry in search_entries]
            
        except Exception as e:
            print(f"❌ Failed to get search history: {e}")
            return []
    
    def get_user_stats(self, user_id: str) -> Dict[str, Any]:
        """Get user statistics"""
        if not self.redis_manager.is_available():
            return {}
        
        client = self.redis_manager.get_client()
        if not client:
            return {}
        
        try:
            stats_key = self._get_user_stats_key(user_id)
            stats = client.hgetall(stats_key)
            
            # Convert numeric values
            if stats:
                stats['total_searches'] = int(stats.get('total_searches', 0))
                stats['avg_response_time'] = float(stats.get('avg_response_time', 0))
                stats['last_results_count'] = int(stats.get('last_results_count', 0))
            
            return stats
            
        except Exception as e:
            print(f"❌ Failed to get user stats: {e}")
            return {}
    
    def save_user_preferences(self, user_id: str, preferences: Dict) -> bool:
        """Save user preferences"""
        if not self.redis_manager.is_available():
            return False
        
        client = self.redis_manager.get_client()
        if not client:
            return False
        
        try:
            prefs_key = self._get_user_preferences_key(user_id)
            client.hset(prefs_key, mapping=preferences)
            client.expire(prefs_key, self.default_ttl)
            return True
            
        except Exception as e:
            print(f"❌ Failed to save preferences: {e}")
            return False
    
    def get_user_preferences(self, user_id: str) -> Dict:
        """Get user preferences"""
        if not self.redis_manager.is_available():
            return {}
        
        client = self.redis_manager.get_client()
        if not client:
            return {}
        
        try:
            prefs_key = self._get_user_preferences_key(user_id)
            return client.hgetall(prefs_key)
            
        except Exception as e:
            print(f"❌ Failed to get preferences: {e}")
            return {}
    
    def get_popular_queries(self, limit: int = 10) -> List[Dict]:
        """Get most popular queries across all users"""
        if not self.redis_manager.is_available():
            return []
        
        client = self.redis_manager.get_client()
        if not client:
            return []
        
        try:
            # Get all user search keys
            pattern = "user_searches:*"
            user_keys = client.keys(pattern)
            
            query_counts = defaultdict(int)
            
            for key in user_keys:
                searches = client.lrange(key, 0, -1)
                for search_json in searches:
                    search_data = json.loads(search_json)
                    query = search_data.get('query', '').lower().strip()
                    if query:
                        query_counts[query] += 1
            
            # Sort by frequency
            popular = sorted(query_counts.items(), key=lambda x: x[1], reverse=True)
            
            return [{"query": query, "count": count} for query, count in popular[:limit]]
            
        except Exception as e:
            print(f"❌ Failed to get popular queries: {e}")
            return []
    
    def clear_user_data(self, user_id: str) -> bool:
        """Clear all data for a specific user"""
        if not self.redis_manager.is_available():
            return False
        
        client = self.redis_manager.get_client()
        if not client:
            return False
        
        try:
            # Delete all user-related keys
            keys_to_delete = [
                self._get_user_searches_key(user_id),
                self._get_user_preferences_key(user_id),
                self._get_user_stats_key(user_id)
            ]
            
            deleted_count = client.delete(*keys_to_delete)
            print(f"✅ Cleared {deleted_count} keys for user: {user_id}")
            return True
            
        except Exception as e:
            print(f"❌ Failed to clear user data: {e}")
            return False

class EnhancedCitationManager:
    """Fixed citation manager with proper clickable citations"""

    def __init__(self):
        self.sources = {}
        self.citation_counter = 1
        self.used_citation_ids = set()
        self.source_id_map = {}
        self.url_to_source_id = {}

    def add_source(self, content: str, metadata: Dict) -> int:
        """Add source with proper deduplication"""
        # Create unique identifier
        url_or_name = metadata.get('url') or metadata.get('file_name', '') or metadata.get('source', '')
        
        # Generate source ID
        if url_or_name:
            source_id = hashlib.md5(url_or_name.encode()).hexdigest()[:16]
        else:
            source_id = hashlib.md5(content[:100].encode()).hexdigest()[:16]
        
        # Check if source already exists
        if source_id in self.sources:
            citation_id = self.sources[source_id]['citation_id']
        else:
            citation_id = self.citation_counter
            self.sources[source_id] = {
                'citation_id': citation_id,
                'metadata': metadata.copy(),
                'content_preview': content[:300] + '...' if len(content) > 300 else content,
                'full_content': content
            }
            self.source_id_map[citation_id] = source_id
            self.citation_counter += 1

        self.used_citation_ids.add(citation_id)
        return citation_id

    def format_inline_citation(self, citation_ids: List[int]) -> str:
        """Format inline citations with proper HTML"""
        if not citation_ids:
            return ""
        
        unique_ids = sorted(set(citation_ids))
        clickable_citations = []
        
        for cid in unique_ids:
            source_id = self.source_id_map.get(cid)
            if source_id and source_id in self.sources:
                data = self.sources[source_id]
                metadata = data['metadata']
                
                if metadata.get('source_type') in ['web_result', 'web_search']:
                    url = metadata.get('url', '')
                    title = metadata.get('title', 'Web Page')
                    if url:
                        clickable_citations.append(
                            f'<a href="{url}" target="_blank" title="{title}" '
                            f'style="text-decoration: none; color: #1f77b4; font-weight: bold;">[{cid}]</a>'
                        )
                    else:
                        clickable_citations.append(f"[{cid}]")
                else:
                    file_name = metadata.get('file_name', metadata.get('source', 'Local Document'))
                    if file_name:
                        file_name = os.path.basename(file_name)
                    clickable_citations.append(
                        f'<span title="{file_name}" style="color: #2ca02c; font-weight: bold;">[{cid}]</span>'
                    )
            else:
                clickable_citations.append(f"[{cid}]")
        
        return "".join(clickable_citations)

    def format_references(self) -> str:
        """Format references section with proper links"""
        if not self.used_citation_ids:
            return ""

        references = []
        for citation_id in sorted(self.used_citation_ids):
            source_id = self.source_id_map.get(citation_id)
            if not source_id or source_id not in self.sources:
                continue

            data = self.sources[source_id]
            metadata = data['metadata']
            
            if metadata.get('source_type') in ['web_result', 'web_search']:
                title = metadata.get('title', 'Web Page')
                url = metadata.get('url', '')
                
                if url:
                    domain = urlparse(url).netloc.replace("www.", "")
                    line = f"[{citation_id}] <a href='{url}' target='_blank'>{title}</a> - {domain}"
                else:
                    line = f"[{citation_id}] {title}"
                    
            elif metadata.get('source_type') == 'ai_summary':
                line = f"[{citation_id}] AI Summary - Tavily Search"
                
            else:
                source_name = metadata.get('file_name', metadata.get('source', 'Unknown Document'))
                if source_name:
                    source_name = os.path.basename(source_name)
                line = f"[{citation_id}] {source_name} (Local Document)"

            # Add preview
            line += f"\n   📄 *Preview: {data['content_preview']}*"
            references.append(line)

        return "\n\n**Sources:**\n" + "\n\n".join(references)

    def clear(self):
        """Clear all citation data"""
        self.sources.clear()
        self.used_citation_ids.clear()
        self.source_id_map.clear()
        self.url_to_source_id.clear()
        self.citation_counter = 1

class EnhancedTavilyWebSearcher:
    """Fixed Tavily web search with proper error handling"""
    
    def __init__(self, cache: UnifiedCache):
        self.cache = cache
        self.tavily_client = None
        
        if WEB_SEARCH_AVAILABLE and TAVILY_API_KEY:
            try:
                self.tavily_client = TavilyClient(api_key=TAVILY_API_KEY)
                print("✅ Tavily client initialized successfully")
            except Exception as e:
                print(f"❌ Tavily client initialization failed: {e}")
        else:
            if not TAVILY_API_KEY:
                print("⚠️ TAVILY_API_KEY not found in environment variables")
        
    def search(self, query: str, max_results: int = 5) -> List[Dict[str, str]]:
        """Enhanced web search with proper error handling"""
        if not self.tavily_client:
            print("❌ Tavily web search not available")
            return []
        
        try:
            # Check cache first
            query_hash = hashlib.md5(query.encode()).hexdigest()[:16]
            cache_key = f"websearch:{query_hash}"
            
            cached_results = self.cache.get(cache_key, "websearch")
            if cached_results:
                print("📋 Using cached web search results")
                return cached_results
            
            print(f"🌐 Searching web for: {query}")
            
            # Perform search with enhanced parameters
            try:
                response = self.tavily_client.search(
                    query=query,
                    search_depth="advanced",
                    max_results=max_results,
                    include_answer=True,
                    include_raw_content=True,
                )
            except Exception as e:
                print(f"❌ Tavily search API error: {e}")
                return []
            
            formatted_results = []
            
            # Process AI summary if available
            if response.get('answer'):
                formatted_results.append({
                    'title': 'AI Summary',
                    'url': 'https://tavily.com',
                    'content': response['answer'],
                    'snippet': response['answer'][:400] + '...' if len(response['answer']) > 400 else response['answer'],
                    'source_type': 'ai_summary',
                    'score': 1.0,
                    'search_query': query
                })
            
            # Process web results
            for result in response.get('results', []):
                content = result.get('content', result.get('raw_content', ''))
                if len(content.strip()) > 50:
                    formatted_result = {
                        'title': result.get('title', 'Web Content'),
                        'url': result.get('url', ''),
                        'content': content,
                        'snippet': content[:500] + '...' if len(content) > 500 else content,
                        'source_type': 'web_result',
                        'score': result.get('score', 0),
                        'search_query': query,
                        'published_date': result.get('published_date', ''),
                        'domain': urlparse(result.get('url', '')).netloc.replace("www.", "")
                    }
                    formatted_results.append(formatted_result)
            
            # Cache results for 1 hour
            if formatted_results:
                self.cache.set(cache_key, formatted_results, ttl=3600, namespace="websearch")
            
            print(f"✅ Found {len(formatted_results)} web results")
            return formatted_results
            
        except Exception as e:
            print(f"❌ Web search failed: {e}")
            return []

class OptimizedDocumentProcessor:
    """Optimized document processor with unified database support"""
    
    SUPPORTED_EXTENSIONS = {
        '.pdf': PyPDFLoader,
        '.txt': TextLoader,
        '.docx': UnstructuredWordDocumentLoader,
        '.csv': CSVLoader,
        '.xlsx': UnstructuredExcelLoader,
        '.md': UnstructuredMarkdownLoader,
        '.json': TextLoader,
        '.html': TextLoader,
        '.xml': TextLoader,
    }
    
    def __init__(self, cache: UnifiedCache, chunk_size: int = 800, chunk_overlap: int = 100):
        self.cache = cache
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ". ", "! ", "? ", ", ", " ", ""],
            length_function=len,
            is_separator_regex=False
        )
    
    def get_file_hash(self, file_path: str) -> str:
        """Fast file hash generation"""
        file_path = Path(file_path)
        stat = file_path.stat()
        return hashlib.sha256(f"{file_path.name}:{stat.st_size}:{stat.st_mtime}".encode()).hexdigest()[:16]
    
    def load_documents_optimized(self, folder_path: str) -> List[Document]:
        """Optimized document loading with parallel processing"""
        folder = Path(folder_path)
        if not folder.exists():
            raise ValueError(f"Folder not found: {folder_path}")
        
        supported_files = [
            str(f) for f in folder.rglob("*") 
            if f.is_file() and f.suffix.lower() in self.SUPPORTED_EXTENSIONS
        ]
        
        if not supported_files:
            print(f"⚠️ No supported documents found in: {folder_path}")
            return []
        
        print(f"📁 Processing {len(supported_files)} files...")
        
        all_documents = []
        
        max_workers = min(8, len(supported_files))
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            future_to_file = {
                executor.submit(self._process_single_file, file_path): file_path
                for file_path in supported_files
            }
            
            for future in as_completed(future_to_file):
                file_path = future_to_file[future]
                try:
                    documents = future.result()
                    all_documents.extend(documents)
                except Exception as e:
                    print(f"❌ Failed to process {file_path}: {e}")
        
        print(f"✅ Processed {len(all_documents)} document chunks")
        return all_documents
    
    def process_web_results(self, web_results: List[Dict]) -> List[Document]:
        """Process web search results into documents with enhanced metadata"""
        if not web_results:
            return []
        
        documents = []
        
        for i, result in enumerate(web_results):
            doc = Document(
                page_content=result.get('content', result.get('snippet', '')),
                metadata={
                    'source': 'web_search',
                    'url': result.get('url', ''),
                    'title': result.get('title', 'Unknown Title'),
                    'source_type': result.get('source_type', 'web_result'),
                    'score': result.get('score', 0),
                    'scraped_at': datetime.now().isoformat(),
                    'type': 'web_result',
                    'search_query': result.get('search_query', ''),
                    'domain': result.get('domain', ''),
                    'published_date': result.get('published_date', '')
                }
            )
            
            url_hash = hashlib.md5(result.get('url', f'web_{i}').encode()).hexdigest()[:8]
            doc.metadata['chunk_id'] = f"web_{url_hash}_{i}"
            
            documents.append(doc)
        
        split_docs = []
        for doc in documents:
            if len(doc.page_content) > self.chunk_size:
                chunks = self.text_splitter.split_documents([doc])
                for j, chunk in enumerate(chunks):
                    chunk.metadata['chunk_id'] = f"{doc.metadata['chunk_id']}_part_{j}"
                split_docs.extend(chunks)
            else:
                split_docs.append(doc)
        
        print(f"✅ Processed {len(split_docs)} web result chunks")
        return split_docs
    
    def _process_single_file(self, file_path: str) -> List[Document]:
        """Process single file with caching"""
        file_hash = self.get_file_hash(file_path)
        cache_key = f"doc:{file_hash}"
        
        cached_docs = self.cache.get(cache_key, "documents")
        if cached_docs:
            return [Document(page_content=doc['content'], metadata=doc['metadata']) 
                   for doc in cached_docs]
        
        try:
            file_ext = Path(file_path).suffix.lower()
            loader_class = self.SUPPORTED_EXTENSIONS[file_ext]
            
            if file_ext == '.csv':
                loader = loader_class(file_path)
            else:
                loader = loader_class(file_path)
            
            documents = loader.load()
            
            for doc in documents:
                doc.metadata.update({
                    'file_name': Path(file_path).name,
                    'file_path': file_path,
                    'processed_at': datetime.now().isoformat(),
                    'file_hash': file_hash,
                    'source_type': 'document'
                })
            
            split_docs = self.text_splitter.split_documents(documents)
            
            for i, doc in enumerate(split_docs):
                doc.metadata['chunk_id'] = f"doc_{file_hash}_{i}"
            
            cacheable_docs = [
                {'content': doc.page_content, 'metadata': doc.metadata}
                for doc in split_docs
            ]
            
            self.cache.set(cache_key, cacheable_docs, ttl=86400, namespace="documents")
            
            return split_docs
            
        except Exception as e:
            print(f"❌ Error processing {file_path}: {e}")
            return []

class OptimizedRAGPipeline:
    """Fixed and optimized RAG pipeline with proper error handling"""
    
    def __init__(self, 
                 docs_folder: str = DEFAULT_DOCS_FOLDER,
                 embedding_model: str = "all-MiniLM-L6-v2",
                 persist_directory: str = "./rag_vectorstore_db",
                 top_k_local: int = 4,
                 top_k_web: int = 4,
                 similarity_threshold: float = 0.3,
                 min_content_length: int = 50):
        
        # Configuration
        self.docs_folder = docs_folder
        self.embedding_model = embedding_model
        self.persist_directory = persist_directory
        self.top_k_local = top_k_local
        self.top_k_web = top_k_web
        self.similarity_threshold = similarity_threshold
        self.min_content_length = min_content_length

        self.vectorstore_client = chromadb.Client(chromadb.config.Settings(
            chroma_db_impl="duckdb+parquet",
            persist_directory=self.persist_directory
        ))
                     
        # Initialize core components
        self.redis_manager = OptimizedRedisManager()
        self.cache = UnifiedCache(self.redis_manager)
        self.citation_manager = EnhancedCitationManager()
        self.web_searcher = EnhancedTavilyWebSearcher(self.cache)
        self.doc_processor = OptimizedDocumentProcessor(self.cache)
        self.user_tracker = UserSearchTracker(self.redis_manager)

        # Lazy-loaded components
        self._embeddings = None
        self._llm = None
        
         # Citation tracking
        self.citations = {}
        self.citation_counter = 0

        # User-specific storage
        self.vectorstores = {}
        self.retriever = None
        
        # Memory storage for conversation histories (fallback)
        self._memory_histories = {}
        
        # Performance optimization
        self.embed_cache = {}
        
        # Initialize default vectorstore
        self._initialize_vectorstore()
        
        # Setup conversation chain
        self._setup_conversation_chain()
        
        print("✅ RAG Pipeline initialized successfully")
    
    @property
    def embeddings(self):
        """Lazy load embeddings"""
        if self._embeddings is None:
            print("🔄 Loading embeddings model...")
            try:
                self._embeddings = HuggingFaceEmbeddings(
                    model_name=self.embedding_model,
                    model_kwargs={'device': 'cpu'},
                    encode_kwargs={'normalize_embeddings': True, 'batch_size': 16}
                )
                print("✅ Embeddings model loaded")
            except Exception as e:
                print(f"❌ Failed to load embeddings: {e}")
                raise
        return self._embeddings
    
    @property
    def llm(self):
        """Lazy load LLM"""
        if self._llm is None:
            print("🔄 Initializing LLM...")
            try:
                self._llm = ChatGroq(
                    api_key=GROQ_API_KEY,
                    model_name="llama3-70b-8192",
                    temperature=0.1,
                    max_tokens=1000,
                    timeout=30,
                    max_retries=2
                )
                print("✅ LLM initialized")
            except Exception as e:
                print(f"❌ Failed to initialize LLM: {e}")
                raise
        return self._llm
    
    def _initialize_vectorstore(self):
        """Initialize default vectorstore"""
        try:
            # Create persist directory if it doesn't exist
            os.makedirs(self.persist_directory, exist_ok=True)
            
            # Try to load existing vectorstore
            if os.path.exists(os.path.join(self.persist_directory, "chroma.sqlite3")):
                print("📂 Loading existing vectorstore...")
                self.default_vectorstore = Chroma(
                    collection_name="default_rag",
                    embedding_function=self.embeddings,
                    persist_directory=self.persist_directory
                )
            else:
                print("🆕 Creating new vectorstore...")
                documents = self.doc_processor.load_documents_optimized(self.docs_folder)
                
                if documents:
                    self.default_vectorstore = Chroma.from_documents(
                        documents,
                        self.embeddings,
                        collection_name="default_rag",
                        persist_directory=self.persist_directory
                    )
                    self.default_vectorstore.persist()
                    print(f"✅ Vectorstore created with {len(documents)} documents")
                else:
                    print("⚠️ No documents found, creating empty vectorstore")
                    self.default_vectorstore = Chroma(
                        collection_name="default_rag",
                        embedding_function=self.embeddings,
                        persist_directory=self.persist_directory
                    )
            
            # Create default retriever
            self.default_retriever = self.default_vectorstore.as_retriever(
                search_type="similarity",
                search_kwargs={"k": self.top_k_local * 2}
            )
            
        except Exception as e:
            print(f"❌ Failed to initialize default vectorstore: {e}")
            self.default_vectorstore = None
            self.default_retriever = None
    
    def _create_vectorstore(self):
        """Create new vectorstore"""
        try:
            # Load documents from folder
            documents = self._load_documents(self.docs_folder)
            if documents:
                self.vectorstore = Chroma.from_documents(
                    documents,
                    self.embeddings,
                    collection_name="unified_rag",
                    persist_directory=self.persist_directory
                )
                self.vectorstore.persist()
        except Exception as e:
            print(f"Failed to create vectorstore: {e}")

    def _load_documents(self, folder_path: str) -> List[Document]:
        """Load .txt, .pdf, and .docx files from the folder into LangChain Document objects."""
        documents = []
        try:
            if os.path.exists(folder_path):
                for filename in os.listdir(folder_path):
                    filepath = os.path.join(folder_path, filename)
                    
                    if filename.endswith('.txt'):
                        try:
                            with open(filepath, 'r', encoding='utf-8') as f:
                                content = f.read()
                        except Exception as e:
                            print(f"Error reading .txt: {filepath} -> {e}")
                            continue

                    elif filename.endswith('.pdf'):
                        try:
                            reader = PdfReader(filepath)
                            content = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
                        except Exception as e:
                            print(f"Error reading .pdf: {filepath} -> {e}")
                            continue

                    elif filename.endswith('.docx'):
                        try:
                            doc = DocxDocument(filepath)
                            content = "\n".join([para.text for para in doc.paragraphs])
                        except Exception as e:
                            print(f"Error reading .docx: {filepath} -> {e}")
                            continue

                    else:
                        continue  # Skip unsupported formats

                    if content and len(content.strip()) > 10:  # Avoid empty docs
                        documents.append(Document(
                            page_content=content,
                            metadata={'source': filepath, 'filename': filename}
                        ))

        except Exception as e:
            print(f"Error loading documents from folder: {e}")
        
        return documents
    
    def _calculate_content_relevance(self, query: str, content: str) -> float:
        """Enhanced relevance calculation optimized for parallel processing"""
        try:
            # Cache common operations to avoid repeated work
            query_lower = query.lower()
            content_lower = content.lower()
            
            # Use sets for faster intersection operations
            query_words = set(query_lower.split())
            content_words = set(content_lower.split())
            
            # Calculate word overlap score
            keyword_overlap = len(query_words.intersection(content_words))
            keyword_score = keyword_overlap / max(len(query_words), 1)
            
            # Calculate phrase matching score
            query_phrases = [phrase.strip() for phrase in query_lower.split() if len(phrase.strip()) > 2]
            phrase_matches = sum(1 for phrase in query_phrases if phrase in content_lower)
            phrase_score = phrase_matches / max(len(query_phrases), 1)
            
            # Calculate content length bonus (prefer substantial chunks)
            length_bonus = min(len(content) / 1000, 0.2)  # Up to 20% bonus for longer content
            
            # Weighted final score
            final_score = (phrase_score * 0.6) + (keyword_score * 0.3) + length_bonus
            return min(final_score, 1.0)
            
        except Exception:
            return 0.0
    
    def _search_local_documents(self, query: str) -> List[Document]:
        """Search local documents with parallel chunk processing"""
        if not self.default_retriever:
            return []
        
        try:
            # Get more candidates for better filtering
            docs = self.default_retriever.invoke(query)
            
            if not docs:
                return []
            
            # **PARALLEL PROCESSING OF DOCUMENT CHUNKS**
            def calculate_relevance_for_doc(doc):
                if len(doc.page_content.strip()) >= self.min_content_length:
                    relevance_score = self._calculate_content_relevance(query, doc.page_content)
                    if relevance_score > 0.1:
                        doc.metadata['relevance_score'] = relevance_score
                        doc.metadata['source_type'] = 'local'
                        return doc
                return None
            
            # Use ThreadPoolExecutor for parallel relevance calculation
            scored_docs = []
            max_workers = min(8, len(docs))  # Limit concurrent threads
            
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                # Submit all relevance calculations in parallel
                future_to_doc = {
                    executor.submit(calculate_relevance_for_doc, doc): doc 
                    for doc in docs
                }
                
                # Collect results as they complete
                for future in as_completed(future_to_doc):
                    try:
                        result = future.result()
                        if result is not None:
                            scored_docs.append(result)
                    except Exception as e:
                        print(f"❌ Error processing document chunk: {e}")
                        continue
            
            # Sort by relevance and return top results
            scored_docs.sort(key=lambda x: x.metadata.get('relevance_score', 0), reverse=True)
            return scored_docs[:self.top_k_local]
            
        except Exception as e:
            print(f"❌ Local search error: {e}")
            return []

    def _search_web(self, query: str) -> List[Document]:
        """Optimized web search using Tavily API with better error handling"""
        try:
            import requests
            
            tavily_api_key = os.getenv("TAVILY_API_KEY")
            if not tavily_api_key:
                return []
            
            # Optimized search parameters
            search_url = "https://api.tavily.com/search"
            headers = {"Authorization": f"Bearer {tavily_api_key}"}
            payload = {
                "query": query,
                "max_results": self.top_k_web + 2,  # Get extra for better filtering
                "include_answer": False,  # Faster response
                "include_raw_content": True,
                "search_depth": "basic"
            }
            
            response = requests.post(search_url, json=payload, headers=headers, timeout=8)
            
            if response.status_code != 200:
                return []
            
            results = response.json().get('results', [])
            
            web_docs = []
            for result in results:
                content = result.get('content', '') or result.get('raw_content', '')
                if len(content.strip()) >= self.min_content_length:
                    
                    # Calculate relevance for web content too
                    relevance_score = self._calculate_content_relevance(query, content)
                    
                    doc = Document(
                        page_content=content,
                        metadata={
                            'source_type': 'web',
                            'title': result.get('title', 'Web Result'),
                            'url': result.get('url', ''),
                            'score': result.get('score', 0),
                            'relevance_score': relevance_score
                        }
                    )
                    web_docs.append(doc)
            
            # Sort by relevance and return top results
            web_docs.sort(key=lambda x: x.metadata.get('relevance_score', 0), reverse=True)
            return web_docs[:self.top_k_web]
            
        except Exception as e:
            return []

    def _add_citation(self, source_info: Dict) -> int:
        """Add citation and return citation number, reusing existing citations for same source"""
        # Create a unique key for the source
        if source_info.get('source_type') == 'local':
            # For local files, use the file path/name as key
            source_key = source_info.get('url', '') or source_info.get('title', '')
            if source_key:
                source_key = os.path.basename(source_key)  # Use just filename
        else:
            # For web sources, use URL as key
            source_key = source_info.get('url', '')
        
        # Check if we already have a citation for this source
        for cite_id, existing_info in self.citations.items():
            existing_key = ''
            if existing_info.get('source_type') == 'local':
                existing_key = existing_info.get('url', '') or existing_info.get('title', '')
                if existing_key:
                    existing_key = os.path.basename(existing_key)
            else:
                existing_key = existing_info.get('url', '')
            
            if source_key and source_key == existing_key:
                return cite_id  # Reuse existing citation number
        
        # If not found, create new citation
        self.citation_counter += 1
        self.citations[self.citation_counter] = {
            'title': source_info.get('title', 'Unknown'),
            'url': source_info.get('url', ''),
            'source_type': source_info.get('source_type', 'web')
        }
        return self.citation_counter

    def _parallel_search(self, query: str) -> List[Document]:
        """Perform parallel search across local and web sources"""
        all_docs = []
        
        with concurrent.futures.ThreadPoolExecutor(max_workers=2) as executor:
            # Submit both searches
            local_future = executor.submit(self._search_local_documents, query)
            web_future = executor.submit(self._search_web, query)
            
            # Wait for completion with timeout
            try:
                local_docs = local_future.result(timeout=10)
                web_docs = web_future.result(timeout=10)
            except concurrent.futures.TimeoutError:
                print("Search timeout - using partial results")
                local_docs = local_future.result() if local_future.done() else []
                web_docs = web_future.result() if web_future.done() else []
        
        # Always try to get both local and web results
        if local_docs:
            all_docs.extend(local_docs)
            
        if web_docs:
            all_docs.extend(web_docs)
        
        # If we have no web results but have local docs, try a simpler web query
        if not web_docs and local_docs:
            try:
                # Try a broader web search
                simplified_query = ' '.join(query.split()[:3])  # First 3 words
                web_docs_fallback = self._search_web(simplified_query)
                if web_docs_fallback:
                    all_docs.extend(web_docs_fallback[:2])  # Add top 2 web results
            except:
                pass
        
        # If we have no local results but have web docs, that's fine too
        if not local_docs and web_docs:
            print(f"No relevant local documents found for query: {query}")
        
        # Remove duplicates and return
        return self._remove_duplicates(all_docs)
    
    def _remove_duplicates(self, documents: List[Document]) -> List[Document]:
        """Remove duplicate documents based on content similarity"""
        if not documents:
            return []
        
        unique_docs = []
        seen_hashes = set()
        
        for doc in documents:
            content_hash = hashlib.md5(doc.page_content[:200].encode()).hexdigest()
            if content_hash not in seen_hashes:
                seen_hashes.add(content_hash)
                unique_docs.append(doc)
        
        return unique_docs
    
    def _format_context_with_citations(self, documents: List[Document], query: str) -> str:
        """Format context with intelligent source prioritization"""
        if not documents:
            return "No relevant information found."
        
        self.citations.clear()
        self.citation_counter = 0
        
        # Separate and score documents by relevance to query
        web_docs = [d for d in documents if d.metadata.get('source_type') == 'web']
        local_docs = [d for d in documents if d.metadata.get('source_type') == 'local']
        
        # Sort all documents by relevance score
        all_docs_scored = []
        for doc in documents:
            relevance = doc.metadata.get('relevance_score', 0)
            all_docs_scored.append((doc, relevance))
        
        # Sort by relevance (highest first)
        all_docs_scored.sort(key=lambda x: x[1], reverse=True)
        
        # Group by source and prioritize most relevant
        source_groups = {}
        for doc, relevance in all_docs_scored:
            if doc.metadata.get('source_type') == 'local':
                source_file = doc.metadata.get('source', 'Unknown')
                source_key = f"local_{os.path.basename(source_file) if source_file else 'Unknown'}"
            else:
                source_key = f"web_{doc.metadata.get('url', f'result_{len(source_groups)}')}"
            
            if source_key not in source_groups:
                source_groups[source_key] = []
            source_groups[source_key].append((doc, relevance))
        
        context_parts = []
        
        # Process each source group, keeping only most relevant content
        for source_key, doc_list in list(source_groups.items())[:6]:  # Limit to top 6 sources for efficiency
            # Get the most relevant document from this source
            doc_list.sort(key=lambda x: x[1], reverse=True)
            best_doc, best_relevance = doc_list[0]
            
            # Skip very low relevance local docs if we have good web content
            if (best_doc.metadata.get('source_type') == 'local' and 
                best_relevance < 0.2 and 
                any(d.metadata.get('source_type') == 'web' and d.metadata.get('relevance_score', 0) > 0.3 for d, _ in all_docs_scored)):
                continue
            
            # Create citation for this source
            if best_doc.metadata.get('source_type') == 'local':
                citation_info = {
                    'title': f"Local Document: {os.path.basename(best_doc.metadata.get('source', 'Unknown'))}",
                    'url': best_doc.metadata.get('source', ''),
                    'source_type': 'local'
                }
            else:
                citation_info = {
                    'title': best_doc.metadata.get('title', 'Web Result'),
                    'url': best_doc.metadata.get('url', ''),
                    'source_type': 'web'
                }
            
            citation_id = self._add_citation(citation_info)
            
            # Combine content from this source (limit length for efficiency)
            combined_content = []
            total_length = 0
            for doc, _ in doc_list:
                if total_length < 800:  # Limit context length per source
                    content_snippet = doc.page_content.strip()
                    combined_content.append(content_snippet)
                    total_length += len(content_snippet)
                else:
                    break
            
            source_label = "(Local PDF)" if citation_info['source_type'] == 'local' else "(Web)"
            context_entry = f"[{citation_id}] {source_label} {' ... '.join(combined_content)}"
            context_parts.append(context_entry)
        
        return "\n\n".join(context_parts)
    
    def _setup_conversation_chain(self):
        """Setup conversation chain with proper history management"""
        def get_session_history(session_id: str):
            """Get session history with Redis fallback"""
            if self.redis_manager.is_available():
                client = self.redis_manager.get_client()
                if client:
                    try:
                        return RedisChatMessageHistory(
                            session_id=session_id,
                            redis_client=client,
                            ttl=86400  # 24 hours
                        )
                    except Exception as e:
                        print(f"⚠️ Redis history failed, using memory: {e}")
            
            # Fallback to in-memory history
            if session_id not in self._memory_histories:
                self._memory_histories[session_id] = InMemoryChatMessageHistory()
            
            return self._memory_histories[session_id]
        
        # Enhanced prompt template
        prompt_template = ChatPromptTemplate.from_messages([
            ("system", """You are an AI assistant specialized exclusively in agriculture.
             

        **CRITICAL DOMAIN CHECK (Execute First):**
        Before doing anything else, determine if the user's question is about:
        - Crops, farming practices, soil, irrigation, fertilizers, pesticides
        - Plant diseases, pest control, livestock, animal husbandry  
        - Agricultural machinery, technology, weather for farming
        - Rural development, agri-markets, agricultural economics

        If the question is NOT about agriculture:
        1. Respond ONLY with: "I'm sorry, I can only assist with agriculture-related questions. Please ask me something about farming, crops, livestock, or related agricultural topics."
        2. Do NOT add any citations, references, or source links
        3. Do NOT search through any provided context
        4. STOP processing immediately

        **FOR AGRICULTURE QUESTIONS ONLY:**
        If and only if the question is agriculture-related:
        - Provide accurate, helpful information about agricultural topics
        - Add citation numbers [1], [2], etc. immediately after factual statements
        - Use exact citation numbers from the context provided
        - Place citations before punctuation: "Fact [1]." not "Fact. [1]"
        - Multiple sources: [1,2] for claims supported by multiple sources
        - Focus on practical, actionable advice when possible
        - Maintain a helpful, professional tone

        **STRICT RULES:**
        - Never provide citations or references in rejection responses
        - Never process context for non-agricultural questions
        - Never explain why a question is outside scope beyond the standard message
        - If unsure whether a question is agricultural, err on the side of rejection

        Examples:
        User: How do I control pests in tomato plants?
        Assistant: [Provide detailed agricultural answer with citations]

        User: Who won the football match yesterday?
        Assistant: I'm sorry, I can only assist with agriculture-related questions. Please ask me something about farming, crops, livestock, or related agricultural topics.

        User: What is the best fertilizer for wheat?
        Assistant: [Provide detailed agricultural answer with citations]"""),

                 MessagesPlaceholder(variable_name="history"),
                ("human", "Context:\n{context}\n\nQuestion: {question}\n\n**CRITICAL INSTRUCTION: Answer ONLY using the provided agricultural context. If this question is not about agriculture, respond with the exact refusal message and provide NO citations whatsoever.**")           ])
        
        # Create chain
        chain = prompt_template | self.llm | StrOutputParser()
        
        self.conversation_chain = RunnableWithMessageHistory(
            chain,
            get_session_history=get_session_history,
            input_messages_key="question",
            history_messages_key="history",
        )
    
    def _format_clickable_citations(self) -> str:
        """Format citations as clickable links"""
        if not self.citations:
            return ""
        
        citations_html = "\n\n**References:**\n"
        for cite_id, info in self.citations.items():
            if info['source_type'] == 'web' and info['url']:
                citations_html += f"[{cite_id}] <a href='{info['url']}' target='_blank'>{info['title']}</a>\n"
            else:
                file_name = os.path.basename(info['title']) if info['title'] else 'Local Document'
                citations_html += f"[{cite_id}] {file_name} (Local PDF)\n"
        
        return citations_html
    
    def _make_citations_clickable(self, response: str) -> str:
        """Convert citation numbers in response to clickable links"""
        import re
        
        def replace_citation(match):
            citation_text = match.group(0)
            citation_nums = re.findall(r'\d+', citation_text)
            
            clickable_parts = []
            for num in citation_nums:
                cite_id = int(num)
                if cite_id in self.citations:
                    citation_info = self.citations[cite_id]
                    if citation_info['source_type'] == 'web' and citation_info['url']:
                        clickable_parts.append(
                            f'<a href="{citation_info["url"]}" target="_blank" title="{citation_info["title"]}" style="text-decoration: none; color: #1f77b4; font-weight: bold;">{num}</a>'
                        )
                    else:
                        file_name = os.path.basename(citation_info['title']) if citation_info['title'] else 'Local Document'
                        clickable_parts.append(
                            f'<span title="{file_name}" style="color: #2ca02c; font-weight: bold;">{num}</span>'
                        )
                else:
                    clickable_parts.append(num)
            
            return f"[{','.join(clickable_parts)}]"
        
        citation_pattern = r'\[(\d+(?:,\s*\d+)*)\]'
        return re.sub(citation_pattern, replace_citation, response)

    def query(self, question: str, user_id: str = "default") -> str:
        """Process query with intelligent content prioritization and clickable citations"""
        start_time = time.time()
        
        translator = get_translator()  # Get the global translator instance
        if translator.is_casual_query(question):
            casual_response = translator.get_casual_response(question, "en")  # Default to English
            response_time = time.time() - start_time
            return f"{casual_response}\n\n*(Response time: {response_time:.2f}s)*"

        try:
            relevant_docs = self._parallel_search(question)
            
            if not relevant_docs:
                broader_terms = question.split()[:4]
                relevant_docs = self._parallel_search(" ".join(broader_terms))
            
            if not relevant_docs:
                return f"I couldn't find relevant information for your query. Please try rephrasing with different keywords."
            
            context = self._format_context_with_citations(relevant_docs, question)
            
            try:
                if self.conversation_chain:
                    response = self.conversation_chain.invoke(
                        {"question": question, "context": context},
                        config={"configurable": {"session_id": f"{user_id}_session"}}
                    )
                else:
                    # Fallback response generation
                    response = self._generate_fallback_response(question, context)
            except Exception as e:
                response = self._generate_fallback_response(question, context)
            
            response_with_clickable_citations = self._make_citations_clickable(response)
            
            unique_local_sources = set()
            unique_web_sources = set()
            
            for doc in relevant_docs:
                if doc.metadata.get('source_type') == 'local':
                    source_file = doc.metadata.get('source', 'Unknown')
                    unique_local_sources.add(os.path.basename(source_file) if source_file else 'Unknown')
                else:
                    unique_web_sources.add(doc.metadata.get('url', 'Unknown'))
            
            local_count = len(unique_local_sources)
            web_count = len(unique_web_sources)
            
            #source_summary = f"\n\n📊 **Sources:** {local_count} local documents, {web_count} web sources"
            citations = self._format_clickable_citations()
            
            response_time = time.time() - start_time
            timing_info = f" (Response time: {response_time:.2f}s)"
            
            return response_with_clickable_citations + timing_info + citations
            
        except Exception as e:
            return f"I encountered an error processing your question. Please try rephrasing your query. Error: {str(e)}"
    
    def _generate_fallback_response(self, question: str, context: str) -> str:
        """Generate fallback response when LLM fails"""
        return f"""Based on the available information regarding: "{question}"

{context}

This information has been compiled from the most relevant sources available. 
Please refer to the cited sources for more detailed information."""
    
    def get_citation_content(self, citation_id: int) -> Optional[Dict]:
        """Get citation information for UI rendering"""
        return self.citations.get(citation_id)
    
    def adjust_similarity_threshold(self, new_threshold: float) -> str:
        """Adjust similarity threshold dynamically"""
        if 0.0 <= new_threshold <= 1.0:
            self.similarity_threshold = new_threshold
            return f"✅ Similarity threshold updated to {new_threshold}"
        return "❌ Threshold must be between 0.0 and 1.0"
    
    def add_documents_for_user(self, user_id: str, documents: List[Document]) -> bool:
        """Add new documents to user's vectorstore"""
        try:
            vectorstore = self.get_user_vectorstore(user_id)
            if not vectorstore:
                print(f"❌ Could not get vectorstore for user: {user_id}")
                return False
            
            # Add documents
            vectorstore.add_documents(documents)
            vectorstore.persist()
            
            print(f"✅ Added {len(documents)} documents for user: {user_id}")
            return True
            
        except Exception as e:
            print(f"❌ Failed to add documents for user {user_id}: {e}")
            return False
    
    def get_system_stats(self) -> Dict[str, Any]:
        """Get system performance statistics"""
        return {
            "vectorstore_available": self.vectorstore is not None,
            "retriever_available": self.retriever is not None,
            "similarity_threshold": self.similarity_threshold,
            "min_content_length": self.min_content_length,
            "citations_count": len(self.citations),
            "cache_size": len(self.embed_cache),
            "optimizations": [
                "Parallel search enabled",
                "Lazy loading implemented", 
                "Reduced filtering strictness",
                "Optimized LLM settings",
                "Citation caching active",
                "Clickable inline citations"
            ]
        }
    
    
    def health_check(self) -> Dict[str, bool]:
        """Comprehensive health check"""
        return {
            "redis_connection": self.redis_manager.is_available(),
            "tavily_api": self.web_searcher.tavily_client,
            "groq_api": GROQ_API_KEY is not None,
            "documents_folder": os.path.exists(self.docs_folder),
            "persist_directory": os.path.exists(self.persist_directory),
            "embeddings_loaded": self._embeddings is not None,
            "llm_loaded": self._llm is not None,
            "default_vectorstore": self.default_vectorstore is not None
        }
    
    def clear_user_data(self, user_id: str) -> bool:
        """Clear all data for a specific user"""
        try:
            # Clear vectorstore
            if user_id in self.vectorstores:
                del self.vectorstores[user_id]
            
            if user_id in self.retriever:
                del self.retriever[user_id]
            
            # Clear memory history
            session_key = f"{user_id}_session"
            if session_key in self._memory_histories:
                del self._memory_histories[session_key]
            
            # Clear Redis history if available
            if self.redis_manager.is_available():
                client = self.redis_manager.get_client()
                if client:
                    try:
                        client.delete(f"message_store:{session_key}")
                    except Exception:
                        pass
            
            print(f"✅ Cleared data for user: {user_id}")
            return True
            
        except Exception as e:
            print(f"❌ Failed to clear data for user {user_id}: {e}")
            return False
        
class InMemoryChatMessageHistory(BaseChatMessageHistory):
    """Simple in-memory chat history for fallback"""
    
    def __init__(self, messages: List = None):
        self.messages = messages or []
    
    def add_message(self, message) -> None:
        self.messages.append(message)
    
    def clear(self) -> None:
        self.messages.clear()

class SessionManager:
    """Fixed session manager with proper Redis integration"""
    
    def __init__(self, redis_manager: OptimizedRedisManager):
        self.redis_manager = redis_manager
        self.sessions_file = "user_sessions.json"
        self.load_sessions()
    
    def load_sessions(self):
        """Load sessions from Redis or fallback to file"""
        try:
            if STREAMLIT_AVAILABLE and not hasattr(st.session_state, 'all_user_sessions'):
                if self.redis_manager.is_available():
                    # Try to load from Redis
                    client = self.redis_manager.get_client()
                    if client:
                        sessions_data = client.get("all_user_sessions")
                        if sessions_data:
                            if STREAMLIT_AVAILABLE:
                                st.session_state.all_user_sessions = json.loads(sessions_data)
                            return
                
                # Fallback to file
                if os.path.exists(self.sessions_file):
                    with open(self.sessions_file, 'r', encoding='utf-8') as f:
                        if STREAMLIT_AVAILABLE:
                            st.session_state.all_user_sessions = json.load(f)
                else:
                    if STREAMLIT_AVAILABLE:
                        st.session_state.all_user_sessions = {}
                        
        except Exception as e:
            print(f"❌ Session loading error: {e}")
            if STREAMLIT_AVAILABLE:
                st.session_state.all_user_sessions = {}
    
    def save_sessions(self):
        """Save sessions to Redis and file"""
        try:
            if not STREAMLIT_AVAILABLE:
                return
                
            sessions_data = st.session_state.get('all_user_sessions', {})
            
            # Save to Redis if available
            if self.redis_manager.is_available():
                client = self.redis_manager.get_client()
                if client:
                    client.setex("all_user_sessions", 86400, json.dumps(sessions_data))
            
            # Also save to file as backup
            with open(self.sessions_file, 'w', encoding='utf-8') as f:
                json.dump(sessions_data, f, ensure_ascii=False, indent=2)
                
        except Exception as e:
            print(f"❌ Session saving error: {e}")
    
    def get_user_sessions(self, user_id: str) -> List[str]:
        """Get all session IDs for a user"""
        if not STREAMLIT_AVAILABLE:
            return []
            
        if user_id not in st.session_state.all_user_sessions:
            st.session_state.all_user_sessions[user_id] = {}
        return list(st.session_state.all_user_sessions[user_id].keys())
    
    def create_new_session(self, user_id: str) -> str:
        """Create a new session for user"""
        if not STREAMLIT_AVAILABLE:
            return f"session_{int(time.time())}"
            
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        session_id = f"session_{int(time.time())}"
        
        if user_id not in st.session_state.all_user_sessions:
            st.session_state.all_user_sessions[user_id] = {}
        
        st.session_state.all_user_sessions[user_id][session_id] = {
            "created_at": timestamp,
            "messages": [],
            "last_updated": timestamp
        }
        self.save_sessions()
        return session_id
    
    def get_session_messages(self, user_id: str, session_id: str) -> List[Dict]:
        """Get messages for a specific session"""
        if not STREAMLIT_AVAILABLE:
            return []
            
        try:
            return st.session_state.all_user_sessions[user_id][session_id]["messages"]
        except KeyError:
            return []
    
    def save_message(self, user_id: str, session_id: str, role: str, content: str):
        """Save a message to session"""
        if not STREAMLIT_AVAILABLE:
            return
            
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        message = {
            "role": role,
            "content": content,
            "timestamp": timestamp
        }
        
        if user_id not in st.session_state.all_user_sessions:
            st.session_state.all_user_sessions[user_id] = {}
        if session_id not in st.session_state.all_user_sessions[user_id]:
            st.session_state.all_user_sessions[user_id][session_id] = {
                "created_at": timestamp,
                "messages": [],
                "last_updated": timestamp
            }
        
        st.session_state.all_user_sessions[user_id][session_id]["messages"].append(message)
        st.session_state.all_user_sessions[user_id][session_id]["last_updated"] = timestamp
        self.save_sessions()

    #-------------------------------------------------------------------------------
class TranslationHelper:
    """Handle multilingual translations"""
    
    def __init__(self):
        try:
            self.translator = Translator()
        except:
            self.translator = None
            print("Warning: Google Translator not available")
            
        self.language_map = {
            "English": "en",
            "Hindi": "hi", 
            "Marathi": "mr",
            "Gujarati": "gu",
            "Tamil": "ta",
            "Telugu": "te",
            "Bengali": "bn",
            "Punjabi": "pa"
        }
    
    # ADD THESE NEW METHODS TO YOUR EXISTING CLASS:
    
    def is_casual_query(self, query: str) -> bool:
        """Detect if query is casual conversation vs information seeking"""
        casual_patterns = [
            r'\b(hi|hello|hey|hola|namaste)\b',
            r'\bhow are you\b',
            r'\bwho are you\b',
            r'\bwhat are you\b',
            r'\bhow can you help\b',
            r'\bwhat can you do\b',
            r'\bgood morning\b',
            r'\bgood evening\b',
            r'\bthanks?\b',
            r'\bthank you\b',
            r'\bbye\b',
            r'\bgoodbye\b',
            r'\bnice to meet you\b',
            r'\bdhanyawad\b',
            r'\bnamaskar\b'
        ]
        
        query_lower = query.lower().strip()
        
        # Check for very short queries (likely casual)
        if len(query_lower.split()) <= 2:
            for pattern in casual_patterns:
                if re.search(pattern, query_lower):
                    return True
        
        # Check for longer casual patterns
        for pattern in casual_patterns:
            if re.search(pattern, query_lower):
                return True
                
        return False
    
    def get_casual_responses(self) -> Dict[str, Dict[str, str]]:
        """Get casual responses in all supported languages"""
        return {
            "en": {
                "greeting": "Hello! I'm Krushi Sarthi, your agricultural assistant. How can I help you with farming or agriculture today?",
                "how_are_you": "I'm doing great, thank you for asking! I'm here to help you with any questions about agriculture, farming, crops, or related topics.",
                "who_are_you": "I'm Krushi Sarthi, an AI-powered agricultural assistant. I can help you with farming techniques, crop management, agricultural best practices, and answer questions using both local documents and web search.",
                "help": "I can assist you with:\n• Farming techniques and best practices\n• Crop diseases and pest management\n• Soil health and fertilizers\n• Weather and seasonal advice\n• Market information\n• Government schemes for farmers\n\nJust ask me any agriculture-related question!",
                "thanks": "You're welcome! Feel free to ask me anything about agriculture or farming.",
                "bye": "Goodbye! Come back anytime you need help with farming or agriculture. Have a great day!",
                "default": "Hello! I'm your agricultural assistant. How can I help you with farming today?"
            },
            "hi": {
                "greeting": "नमस्ते! मैं कृषि सारथी हूं, आपका कृषि सहायक। आज मैं खेती या कृषि में आपकी कैसे मदद कर सकता हूं?",
                "how_are_you": "मैं बहुत अच्छा हूं, पूछने के लिए धन्यवाद! मैं यहां कृषि, खेती, फसलों के बारे में आपके किसी भी सवाल में मदद करने के लिए हूं।",
                "who_are_you": "मैं कृषि सारथी हूं, एक AI-संचालित कृषि सहायक। मैं खेती की तकनीकों, फसल प्रबंधन, कृषि की बेहतरीन प्रथाओं में मदद कर सकता हूं।",
                "help": "मैं इन चीजों में आपकी मदद कर सकता हूं:\n• खेती की तकनीकें और बेस्ट प्रैक्टिसेज\n• फसल की बीमारियां और कीट प्रबंधन\n• मिट्टी का स्वास्थ्य और उर्वरक\n• मौसम और मौसमी सलाह\n• बाजार की जानकारी\n• किसानों के लिए सरकारी योजनाएं\n\nकृषि से जुड़ा कोई भी सवाल पूछें!",
                "thanks": "आपका स्वागत है! कृषि या खेती के बारे में कुछ भी पूछने में संकोच न करें।",
                "bye": "अलविदा! खेती या कृषि में मदद की जरूरत हो तो कभी भी आएं। आपका दिन शुभ हो!",
                "default": "नमस्ते! मैं आपका कृषि सहायक हूं। आज खेती में कैसे मदद कर सकता हूं?"
            },
            "mr": {
                "greeting": "नमस्कार! मी कृषी सारथी आहे, तुमचा कृषी सहाय्यक। आज मी तुम्हाला शेती किंवा कृषीमध्ये कशी मदत करू शकतो?",
                "how_are_you": "मी खूप चांगला आहे, विचारल्याबद्दल धन्यवाद! मी येथे कृषी, शेती, पिकांविषयी तुमच्या कोणत्याही प्रश्नात मदत करण्यासाठी आहे.",
                "who_are_you": "मी कृषी सारथी आहे, एक AI-चालित कृषी सहाय्यक. मी शेतीच्या तंत्रांमध्ये, पीक व्यवस्थापनात, कृषी सराव मध्ये मदत करू शकतो.",
                "help": "मी या गोष्टींमध्ये तुमची मदत करू शकतो:\n• शेतीची तंत्रे आणि चांगल्या पद्धती\n• पिकांचे आजार आणि कीटक व्यवस्थापन\n• मातीचे आरोग्य आणि खते\n• हवामान आणि हंगामी सल्ला\n• बाजाराची माहिती\n• शेतकऱ्यांसाठी सरकारी योजना\n\nकृषीशी संबंधित कोणताही प्रश्न विचारा!",
                "thanks": "तुमचे स्वागत आहे! कृषी किंवा शेतीविषयी काहीही विचारण्यास मोकळ्या मनाने या.",
                "bye": "निरोप! शेती किंवा कृषीमध्ये मदत हवी असेल तर कधीही या. तुमचा दिवस चांगला जावो!",
                "default": "नमस्कार! मी तुमचा कृषी सहाय्यक आहे. आज शेतीमध्ये कशी मदत करू शकतो?"
            },
            "gu": {
                "greeting": "નમસ્તે! હું કૃષિ સારથી છું, તમારો કૃષિ સહાયક. આજે હું તમને ખેતી અથવા કૃષિમાં કેવી રીતે મદદ કરી શકું?",
                "how_are_you": "હું ખૂબ સારો છું, પૂછવા બદલ આભાર! હું અહીં કૃષિ, ખેતી, પાકો વિશેના તમારા કોઈપણ પ્રશ્નોમાં મદદ કરવા માટે છું.",
                "who_are_you": "હું કૃષિ સારથી છું, એક AI-સંચાલિત કૃષિ સહાયક. હું ખેતીની તકનીકો, પાક વ્યવસ્થાપન, કૃષિ શ્રેષ્ઠ પ્રથાઓમાં મદદ કરી શકું છું.",
                "help": "હું આ બાબતોમાં તમારી મદદ કરી શકું છું:\n• ખેતીની તકનીકો અને શ્રેષ્ઠ પ્રથાઓ\n• પાકના રોગો અને જીવાત વ્યવસ્થાપન\n• માટીનું સ્વાસ્થ્ય અને ખાતર\n• હવામાન અને મોસમી સલાહ\n• બજારની માહિતી\n• ખેડૂતો માટે સરકારી યોજનાઓ\n\nકૃષિ સંબંધિત કોઈપણ પ્રશ્ન પૂછો!",
                "thanks": "તમારું સ્વાગત છે! કૃષિ અથવા ખેતી વિશે કંઈપણ પૂછવામાં સંકોચ ન કરો.",
                "bye": "વિદાય! ખેતી અથવા કૃષિમાં મદદની જરૂર હોય તો કોઈ પણ સમયે આવો. તમારો દિવસ સારો જાય!",
                "default": "નમસ્તે! હું તમારો કૃષિ સહાયક છું. આજે ખેતીમાં કેવી રીતે મદદ કરી શકું?"
            },
            "ta": {
                "greeting": "வணக்கம்! நான் கிருஷி சாரதி, உங்கள் விவசாய உதவியாளர். இன்று நான் விவசாயம் அல்லது வேளாண்மையில் உங்களுக்கு எப்படி உதவ முடியும்?",
                "how_are_you": "நான் நன்றாக இருக்கிறேன், கேட்டதற்கு நன்றி! விவசாயம், வேளாண்மை, பயிர்கள் பற்றிய உங்கள் எந்த கேள்விகளுக்கும் நான் இங்கே உதவ இருக்கிறேன்.",
                "who_are_you": "நான் கிருஷி சாரதி, AI-இயங்கும் விவசாய உதவியாளர். விவசாய நுட்பங்கள், பயிர் மேலாண்மை, வேளாண் சிறந்த நடைமுறைகளில் நான் உதவ முடியும்.",
                "help": "நான் இந்த விஷயங்களில் உங்களுக்கு உதவ முடியும்:\n• விவசாய நுட்பங்கள் மற்றும் சிறந்த நடைமுறைகள்\n• பயிர் நோய்கள் மற்றும் பூச்சி மேலாண்மை\n• மண் ஆரோக்கியம் மற்றும் உரங்கள்\n• வானிலை மற்றும் பருவகால ஆலோசனை\n• சந்தை தகவல்\n• விவசாயிகளுக்கான அரசு திட்டங்கள்\n\nவேளாண்மை தொடர்பான எந்த கேள்வியும் கேளுங்கள்!",
                "thanks": "வரவேற்கிறோம்! விவசாயம் அல்லது வேளாண்மை பற்றி எதையும் கேட்க தயங்க வேண்டாம்.",
                "bye": "வணக்கம்! விவசாயம் அல்லது வேளாண்மையில் உதவி தேவைப்பட்டால் எப்போது வேண்டுமானாலும் வாருங்கள். உங்கள் நாள் நல்லதாக இருக்கட்டும்!",
                "default": "வணக்கம்! நான் உங்கள் விவசாய உதவியாளர். இன்று விவசாயத்தில் எப்படி உதவ முடியும்?"
            },
            "te": {
                "greeting": "నమస్కారం! నేను కృషి సారథి, మీ వ్యవసాయ సహాయకుడను. ఈ రోజు వ్యవసాయం లేదా కృషిలో నేను మీకు ఎలా సహాయం చేయగలను?",
                "how_are_you": "నేను బాగున్నాను, అడిగినందుకు ధన్యవాదాలు! వ్యవసాయం, కృషి, పంటల గురించి మీ ఏ ప్రశ్నలకైనా సహాయం చేయడానికి నేను ఇక్కడ ఉన్నాను.",
                "who_are_you": "నేను కృషి సారథి, AI-ఆధారిత వ్యవసాయ సహాయకుడను. వ్యవసాయ పద్ధతులు, పంట నిర్వహణ, వ్యవసాయ ఉత్తమ అభ్యాసాలలో నేను సహాయం చేయగలను.",
                "help": "నేను ఈ విషయాలలో మీకు సహాయం చేయగలను:\n• వ్యవసాయ పద్ధతులు మరియు ఉత్తమ అభ్యాసాలు\n• పంట వ్యాధులు మరియు కీట నిర్వహణ\n• మట్టి ఆరోగ్యం మరియు ఎరువులు\n• వాతావరణం మరియు కాలానుగుణ సలహా\n• మార్కెట్ సమాచారం\n• రైతుల కోసం ప్రభుత్వ పథకాలు\n\nవ్యవసాయ సంబంధిత ఏ ప్రశ్న అయినా అడగండి!",
                "thanks": "మీకు స్వాగతం! వ్యవసాయం లేదా కృషి గురించి ఏదైనా అడగడానికి సంకోచించకండి.",
                "bye": "నమస్కారం! వ్యవసాయం లేదా కృషిలో సహాయం అవసరమైతే ఎప్పుడైనా రండి. మీ రోజు మంచిగా గడవాలి!",
                "default": "నమస్కారం! నేను మీ వ్యవసాయ సహాయకుడను. ఈ రోజు వ్యవసాయంలో ఎలా సహాయం చేయగలను?"
            },
            "bn": {
                "greeting": "নমস্কার! আমি কৃষি সারথি, আপনার কৃষি সহায়ক। আজ আমি কৃষি বা কৃষিকাজে আপনাকে কীভাবে সাহায্য করতে পারি?",
                "how_are_you": "আমি খুবই ভালো আছি, জিজ্ঞাসা করার জন্য ধন্যবাদ! কৃষি, চাষাবাদ, ফসল সম্পর্কে আপনার যেকোনো প্রশ্নে সাহায্য করার জন্য আমি এখানে আছি।",
                "who_are_you": "আমি কৃষি সারথি, একটি AI-চালিত কৃষি সহায়ক। আমি কৃষি কৌশল, ফসল ব্যবস্থাপনা, কৃষি সর্বোত্তম অনুশীলনে সাহায্য করতে পারি।",
                "help": "আমি এই বিষয়গুলিতে আপনাকে সাহায্য করতে পারি:\n• কৃষি কৌশল এবং সর্বোত্তম অনুশীলন\n• ফসলের রোগ এবং পোকা ব্যবস্থাপনা\n• মাটির স্বাস্থ্য এবং সার\n• আবহাওয়া এবং মৌসুমী পরামর্শ\n• বাজারের তথ্য\n• কৃষকদের জন্য সরকারি প্রকল্প\n\nকৃষি সংক্রান্ত যেকোনো প্রশ্ন করুন!",
                "thanks": "আপনাকে স্বাগতম! কৃষি বা চাষাবাদ সম্পর্কে যেকোনো কিছু জিজ্ঞাসা করতে দ্বিধা করবেন না।",
                "bye": "বিদায়! কৃষি বা চাষাবাদে সাহায্যের প্রয়োজন হলে যেকোনো সময় আসুন। আপনার দিনটি ভালো কাটুক!",
                "default": "নমস্কার! আমি আপনার কৃষি সহায়ক। আজ কৃষিকাজে কীভাবে সাহায্য করতে পারি?"
            },
            "pa": {
                "greeting": "ਸਤ ਸ੍ਰੀ ਅਕਾਲ! ਮੈਂ ਕ੍ਰਿਸ਼ੀ ਸਾਰਥੀ ਹਾਂ, ਤੁਹਾਡਾ ਖੇਤੀਬਾੜੀ ਸਹਾਇਕ। ਅੱਜ ਮੈਂ ਖੇਤੀ ਜਾਂ ਖੇਤੀਬਾੜੀ ਵਿੱਚ ਤੁਹਾਡੀ ਕਿਵੇਂ ਮਦਦ ਕਰ ਸਕਦਾ ਹਾਂ?",
                "how_are_you": "ਮੈਂ ਬਹੁਤ ਚੰਗਾ ਹਾਂ, ਪੁੱਛਣ ਲਈ ਧੰਨਵਾਦ! ਮੈਂ ਇੱਥੇ ਖੇਤੀਬਾੜੀ, ਖੇਤੀ, ਫਸਲਾਂ ਬਾਰੇ ਤੁਹਾਡੇ ਕਿਸੇ ਵੀ ਸਵਾਲ ਦੀ ਮਦਦ ਲਈ ਹਾਂ।",
                "who_are_you": "ਮੈਂ ਕ੍ਰਿਸ਼ੀ ਸਾਰਥੀ ਹਾਂ, ਇੱਕ AI-ਸੰਚਾਲਿਤ ਖੇਤੀਬਾੜੀ ਸਹਾਇਕ। ਮੈਂ ਖੇਤੀ ਦੀਆਂ ਤਕਨੀਕਾਂ, ਫਸਲ ਪ੍ਰਬੰਧਨ, ਖੇਤੀਬਾੜੀ ਦੀਆਂ ਵਧੀਆ ਪ੍ਰਥਾਵਾਂ ਵਿੱਚ ਮਦਦ ਕਰ ਸਕਦਾ ਹਾਂ।",
                "help": "ਮੈਂ ਇਹਨਾਂ ਗੱਲਾਂ ਵਿੱਚ ਤੁਹਾਡੀ ਮਦਦ ਕਰ ਸਕਦਾ ਹਾਂ:\n• ਖੇਤੀ ਦੀਆਂ ਤਕਨੀਕਾਂ ਅਤੇ ਵਧੀਆ ਪ੍ਰਥਾਵਾਂ\n• ਫਸਲਾਂ ਦੀਆਂ ਬਿਮਾਰੀਆਂ ਅਤੇ ਕੀੜੇ-ਮਕੌੜੇ ਪ੍ਰਬੰਧਨ\n• ਮਿੱਟੀ ਦੀ ਸਿਹਤ ਅਤੇ ਖਾਦ\n• ਮੌਸਮ ਅਤੇ ਮੌਸਮੀ ਸਲਾਹ\n• ਮਾਰਕੀਟ ਦੀ ਜਾਣਕਾਰੀ\n• ਕਿਸਾਨਾਂ ਲਈ ਸਰਕਾਰੀ ਸਕੀਮਾਂ\n\nਖੇਤੀਬਾੜੀ ਨਾਲ ਜੁੜਿਆ ਕੋਈ ਵੀ ਸਵਾਲ ਪੁੱਛੋ!",
                "thanks": "ਤੁਹਾਡਾ ਸਵਾਗਤ ਹੈ! ਖੇਤੀਬਾੜੀ ਜਾਂ ਖੇਤੀ ਬਾਰੇ ਕੁਝ ਵੀ ਪੁੱਛਣ ਵਿੱਚ ਸੰਕੋਚ ਨਾ ਕਰੋ।",
                "bye": "ਸਤ ਸ੍ਰੀ ਅਕਾਲ! ਖੇਤੀਬਾੜੀ ਜਾਂ ਖੇਤੀ ਵਿੱਚ ਮਦਦ ਦੀ ਲੋੜ ਹੋਵੇ ਤਾਂ ਕਿਸੇ ਵੀ ਸਮੇਂ ਆਓ। ਤੁਹਾਡਾ ਦਿਨ ਚੰਗਾ ਰਹੇ!",
                "default": "ਸਤ ਸ੍ਰੀ ਅਕਾਲ! ਮੈਂ ਤੁਹਾਡਾ ਖੇਤੀਬਾੜੀ ਸਹਾਇਕ ਹਾਂ। ਅੱਜ ਖੇਤੀ ਵਿੱਚ ਕਿਵੇਂ ਮਦਦ ਕਰ ਸਕਦਾ ਹਾਂ?"
            }
        }
    
    def get_casual_response(self, query: str, lang_code: str = "en") -> str:
        """Generate appropriate casual responses based on query and language"""
        query_lower = query.lower().strip()
        responses = self.get_casual_responses()
        lang_responses = responses.get(lang_code, responses["en"])
        
        # Determine response type based on query content
        if any(greeting in query_lower for greeting in ["hi", "hello", "hey", "namaste", "namaskar", "sat sri akal"]):
            return lang_responses["greeting"]
        elif "how are you" in query_lower or "कैसे हो" in query_lower or "कसे आहात" in query_lower:
            return lang_responses["how_are_you"]
        elif any(phrase in query_lower for phrase in ["who are you", "what are you", "तुम कौन हो", "तुम्ही कोण आहात"]):
            return lang_responses["who_are_you"]
        elif any(phrase in query_lower for phrase in ["help", "can you do", "मदद", "સહાય", "உதவி"]):
            return lang_responses["help"]
        elif any(phrase in query_lower for phrase in ["thank", "thanks", "धन्यवाद", "धन्यावाद"]):
            return lang_responses["thanks"]
        elif any(phrase in query_lower for phrase in ["bye", "goodbye", "अलविदा", "निरोप"]):
            return lang_responses["bye"]
        else:
            return lang_responses["default"]

    # KEEP ALL YOUR EXISTING METHODS (translate_text, get_ui_text) AS THEY ARE
    
    def translate_text(self, text: str, target_lang: str) -> str:
        """Translate text to target language"""
        if target_lang == "en" or not self.translator:
            return text
        
        try:
            if '<a href=' in text or '[' in text and ']' in text:
                parts = text.split('\n\n**References:**')
                if len(parts) == 2:
                    main_content = parts[0]
                    references = '\n\n**References:**' + parts[1]
                    translated_main = self.translator.translate(main_content, dest=target_lang).text
                    return translated_main + references
            
            result = self.translator.translate(text, dest=target_lang)
            return result.text
        except Exception as e:
            print(f"Translation error: {e}")
            return text
    
    def get_ui_text(self, key: str, lang_code: str) -> str:
        """Get UI text in selected language"""
        ui_texts = {
            "en": {
                "title": "🌾 Krushi Sarthi Chatbot",
                "user_id_prompt": "Enter your User ID:",
                "language_select": "Select Language",
                "new_session": "New Session",
                "previous_sessions": "Previous Sessions", 
                "chat_input": "What is your query?",
                "searching": "Searching documents and web...",
                "no_sessions": "No previous sessions found",
                "session_created": "New session created successfully!",
                "welcome_message": "Welcome! I'm your agricultural assistant. Ask me anything about farming, crops, or agriculture."
            },
            "hi": {
                "title": "🌾 कृषि सारथी चैटबॉट",
                "user_id_prompt": "अपना उपयोगकर्ता ID दर्ज करें:",
                "language_select": "भाषा चुनें",
                "new_session": "नया सत्र",
                "previous_sessions": "पिछले सत्र",
                "chat_input": "आपका प्रश्न क्या है?",
                "searching": "दस्तावेज़ और वेब खोज रहे हैं...",
                "no_sessions": "कोई पिछला सत्र नहीं मिला",
                "session_created": "नया सत्र सफलतापूर्वक बनाया गया!",
                "welcome_message": "नमस्ते! मैं आपका कृषि सहायक हूं। खेती, फसलों या कृषि के बारे में कुछ भी पूछें।"
            },
            "mr": {
                "title": "🌾 कृषी सारथी चॅटबॉट",
                "user_id_prompt": "तुमचा वापरकर्ता ID टाका:",
                "language_select": "भाषा निवडा",
                "new_session": "नवीन सत्र",
                "previous_sessions": "मागील सत्रे",
                "chat_input": "तुमचा प्रश्न काय आहे?",
                "searching": "कागदपत्रे आणि वेब शोधत आहे...",
                "no_sessions": "कोणतेही मागील सत्र सापडले नाही",
                "session_created": "नवीन सत्र यशस्वीरित्या तयार केले!",
                "welcome_message": "नमस्कार! मी तुमचा कृषी सहाय्यक आहे। शेती, पिके किंवा कृषी संबंधी काहीही विचारा।"
            },
            "gu": {
                "title": "🌾 કૃષિ સારથી ચેટબોટ",
                "user_id_prompt": "તમારો વપરાશકર્તા ID દાખલ કરો:",
                "language_select": "ભાષા પસંદ કરો",
                "new_session": "નવું સત્ર",
                "previous_sessions": "અગાઉના સત્રો",
                "chat_input": "તમારો પ્રશ્ન શું છે?",
                "searching": "દસ્તાવેજો અને વેબ શોધી રહ્યા છીએ...",
                "no_sessions": "કોઈ અગાઉના સત્રો મળ્યા નથી",
                "session_created": "નવું સત્ર સફળતાપૂર્વક બનાવવામાં આવ્યું!",
                "welcome_message": "સ્વાગત છે! હું તમારો કૃષિ સહાયક છું। ખેતી, પાકો અથવા કૃષિ વિશે કંઈપણ પૂછો।"
            },
            "ta": {
                "title": "🌾 கிருஷி சாரதி சாட்போட்",
                "user_id_prompt": "உங்கள் பயனர் ID ஐ உள்ளிடவும்:",
                "language_select": "மொழியைத் தேர்ந்தெடுக்கவும்",
                "new_session": "புதிய அமர்வு",
                "previous_sessions": "முந்தைய அமர்வுகள்",
                "chat_input": "உங்கள் கேள்வி என்ன?",
                "searching": "ஆவணங்கள் மற்றும் வலையைத் தேடுகிறது...",
                "no_sessions": "முந்தைய அமர்வுகள் எதுவும் கிடைக்கவில்லை",
                "session_created": "புதிய அமர்வு வெற்றிகரமாக உருவாக்கப்பட்டது!",
                "welcome_message": "வரவேற்கிறோம்! நான் உங்கள் விவசாய உதவியாளர். விவசாயம், பயிர்கள் அல்லது வேளாண்மை பற்றி எதையும் கேளுங்கள்."
            },
            "te": {
                "title": "🌾 కృషి సారథి చాట్‌బాట్",
                "user_id_prompt": "మీ వినియోగదారు IDని నమోదు చేయండి:",
                "language_select": "భాషను ఎంచుకోండి",
                "new_session": "కొత్త సెషన్",
                "previous_sessions": "మునుపటి సెషన్లు",
                "chat_input": "మీ ప్రశ్న ఏమిటి?",
                "searching": "పత్రాలు మరియు వెబ్‌ను వెతుకుతోంది...",
                "no_sessions": "మునుపటి సెషన్లు ఏవీ కనుగొనబడలేదు",
                "session_created": "కొత్త సెషన్ విజయవంతంగా సృష్టించబడింది!",
                "welcome_message": "స్వాగతం! నేను మీ వ్యవసాయ సహాయకుడిని. వ్యవసాయం, పంటలు లేదా వ్యవసాయం గురించి ఏదైనా అడగండి."
            },
            "bn": {
                "title": "🌾 কৃষি সারথি চ্যাটবট",
                "user_id_prompt": "আপনার ব্যবহারকারী ID প্রবেশ করান:",
                "language_select": "ভাষা নির্বাচন করুন",
                "new_session": "নতুন সেশন",
                "previous_sessions": "পূর্ববর্তী সেশনগুলি",
                "chat_input": "আপনার প্রশ্ন কি?",
                "searching": "নথি এবং ওয়েব অনুসন্ধান করছে...",
                "no_sessions": "কোনো পূর্ববর্তী সেশন পাওয়া যায়নি",
                "session_created": "নতুন সেশন সফলভাবে তৈরি হয়েছে!",
                "welcome_message": "স্বাগতম! আমি আপনার কৃষি সহায়ক। কৃষি, ফসল বা কৃষি সম্পর্কে যে কোনো কিছু জিজ্ঞাসা করুন।"
            },
            "pa": {
                "title": "🌾 ਕ੍ਰਿਸ਼ੀ ਸਾਰਥੀ ਚੈਟਬੋਟ",
                "user_id_prompt": "ਆਪਣਾ ਯੂਜ਼ਰ ID ਦਾਖਲ ਕਰੋ:",
                "language_select": "ਭਾਸ਼ਾ ਚੁਣੋ",
                "new_session": "ਨਵਾਂ ਸੈਸ਼ਨ",
                "previous_sessions": "ਪਿਛਲੇ ਸੈਸ਼ਨ",
                "chat_input": "ਤੁਹਾਡਾ ਸਵਾਲ ਕੀ ਹੈ?",
                "searching": "ਦਸਤਾਵੇਜ਼ ਅਤੇ ਵੈੱਬ ਖੋਜ ਰਿਹਾ ਹੈ...",
                "no_sessions": "ਕੋਈ ਪਿਛਲੇ ਸੈਸ਼ਨ ਨਹੀਂ ਮਿਲੇ",
                "session_created": "ਨਵਾਂ ਸੈਸ਼ਨ ਸਫਲਤਾਪੂਰਵਕ ਬਣਾਇਆ ਗਿਆ!",
                "welcome_message": "ਜੀ ਆਇਆਂ ਨੂੰ! ਮੈਂ ਤੁਹਾਡਾ ਖੇਤੀਬਾੜੀ ਸਹਾਇਕ ਹਾਂ। ਖੇਤੀ, ਫਸਲਾਂ ਜਾਂ ਖੇਤੀਬਾੜੀ ਬਾਰੇ ਕੁਝ ਵੀ ਪੁੱਛੋ।"
            }
        }
        return ui_texts.get(lang_code, ui_texts["en"]).get(key, key)

# Global instances (Fixed)
_pipeline_instance = None
_session_manager = None
_translator = None

def get_pipeline_instance():
    """Get or create global pipeline instance"""
    global _pipeline_instance
    if _pipeline_instance is None:
        print("🚀 Initializing RAG Pipeline...")
        _pipeline_instance = OptimizedRAGPipeline()
    return _pipeline_instance

def get_session_manager():
    """Get or create session manager"""
    global _session_manager
    if _session_manager is None:
        pipeline = get_pipeline_instance()
        _session_manager = SessionManager(pipeline.redis_manager)
    return _session_manager

def get_translator():
    """Get or create translator"""
    global _translator
    if _translator is None:
        _translator = TranslationHelper()
    return _translator

def enhanced_chatbot_interface():
    """Enhanced Streamlit chatbot interface with multilingual support and session management"""
    
    # Initialize session manager and translator
    session_manager = get_session_manager()
    translator = get_translator()
    
    # Initialize session state variables
    if "current_user_id" not in st.session_state:
        st.session_state.current_user_id = None
    if "current_session_id" not in st.session_state:
        st.session_state.current_session_id = None
    if "selected_language" not in st.session_state:
        st.session_state.selected_language = "English"
    if "all_user_sessions" not in st.session_state:
        st.session_state.all_user_sessions = {}
    
    # Language selection in sidebar
    st.sidebar.title("🌐 Settings")
    
    # Language selector
    available_languages = list(translator.language_map.keys())
    selected_language = st.sidebar.selectbox(
        translator.get_ui_text("language_select", translator.language_map[st.session_state.selected_language]),
        available_languages,
        index=available_languages.index(st.session_state.selected_language)
    )
    
    if selected_language != st.session_state.selected_language:
        st.session_state.selected_language = selected_language
        st.rerun()
    
    lang_code = translator.language_map[selected_language]
    
    # User authentication
    if st.session_state.current_user_id is None:
        st.title(translator.get_ui_text("title", lang_code))
        
        # User ID input
        user_id_input = st.text_input(
            translator.get_ui_text("user_id_prompt", lang_code),
            placeholder="Enter your unique ID (e.g., farmer123)"
        )
        
        if st.button("Start Session", type="primary"):
            if user_id_input.strip():
                st.session_state.current_user_id = user_id_input.strip()
                # Create a new session by default
                st.session_state.current_session_id = session_manager.create_new_session(st.session_state.current_user_id)
                st.success(translator.get_ui_text("session_created", lang_code))
                st.rerun()
            else:
                st.error("Please enter a valid User ID")
        
        # Show welcome message
        st.info(translator.get_ui_text("welcome_message", lang_code))
        return
    
    # Main interface for authenticated users
    st.title(translator.get_ui_text("title", lang_code))
    st.sidebar.markdown(f"**User:** {st.session_state.current_user_id}")
    
    # Session management in sidebar
    st.sidebar.markdown("---")
    st.sidebar.subheader("📁 Session Management")
    
    # New session button
    if st.sidebar.button(translator.get_ui_text("new_session", lang_code)):
        st.session_state.current_session_id = session_manager.create_new_session(st.session_state.current_user_id)
        st.sidebar.success(translator.get_ui_text("session_created", lang_code))
        st.rerun()
    
    # Previous sessions
    user_sessions = session_manager.get_user_sessions(st.session_state.current_user_id)
    if user_sessions:
        st.sidebar.subheader(translator.get_ui_text("previous_sessions", lang_code))
        
        # Sort sessions by creation time (newest first)
        sorted_sessions = []
        for session_id in user_sessions:
            try:
                session_data = st.session_state.all_user_sessions[st.session_state.current_user_id][session_id]
                sorted_sessions.append((session_id, session_data.get("created_at", "")))
            except KeyError:
                continue
        
        sorted_sessions.sort(key=lambda x: x[1], reverse=True)
        
        for session_id, created_at in sorted_sessions[:10]:  # Show last 10 sessions
            session_label = f"Session {created_at[:16]}" if created_at else session_id
            if st.sidebar.button(session_label, key=f"session_{session_id}"):
                st.session_state.current_session_id = session_id
                st.rerun()
    else:
        st.sidebar.info(translator.get_ui_text("no_sessions", lang_code))
    
    # Logout button
    if st.sidebar.button("🚪 Logout"):
        st.session_state.current_user_id = None
        st.session_state.current_session_id = None
        st.rerun()
    
    # Load and display chat history for current session
    if st.session_state.current_session_id:
        messages = session_manager.get_session_messages(
            st.session_state.current_user_id, 
            st.session_state.current_session_id
        )
        
        # Display chat messages
        for message in messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"], unsafe_allow_html=True)
        
        # Chat input
        if prompt := st.chat_input(translator.get_ui_text("chat_input", lang_code)):
            # Save user message
            session_manager.save_message(
                st.session_state.current_user_id,
                st.session_state.current_session_id,
                "user",
                prompt
            )
            
            # Display user message
            with st.chat_message("user"):
                st.markdown(prompt)
            
            # Get assistant response
            with st.chat_message("assistant"):
                response = ""

                if translator.is_casual_query(prompt):
                    response = translator.get_casual_response(prompt, lang_code)
                else:
                    with st.spinner(translator.get_ui_text("searching", lang_code)):                 
                        english_query = prompt
                        if lang_code != "en" and translator.translator:
                            try:
                                english_query = translator.translator.translate(prompt, dest="en").text
                            except:
                             english_query = prompt
                    
                    # Get response from pipeline
                    pipeline = get_pipeline_instance()
                    response = pipeline.query(english_query, st.session_state.current_user_id)
                    
                    # Translate response if needed
                    if lang_code != "en":
                        response = translator.translate_text(response, lang_code)
                
                # Display response with typing effect
                response_placeholder = st.empty()
                full_response = ""
                
                # Simple typing effect
                words = response.split(" ")
                for i, word in enumerate(words):
                    full_response += word + " "
                    if i % 10 == 0:  # Update every 10 words for smoother effect
                        response_placeholder.markdown(full_response + "▌", unsafe_allow_html=True)
                        time.sleep(0.05)
                
                response_placeholder.markdown(response, unsafe_allow_html=True)
                
                # Save assistant response
                session_manager.save_message(
                    st.session_state.current_user_id,
                    st.session_state.current_session_id,
                    "assistant",
                    response
                )

def main():
    """Main function with enhanced error handling"""
    if not STREAMLIT_AVAILABLE:
        print("❌ Streamlit not available. Running in CLI mode...")
        # Add CLI interface here if needed
        return
    
    # Configure Streamlit page
    st.set_page_config(
        page_title="Krushi Sarthi Chatbot",
        page_icon="🌾",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Enhanced Custom CSS for better visibility
    st.markdown("""
    <style>
    /* Import Google Fonts for better readability */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
    
    /* Root variables for consistent theming */
    :root {
        --primary-bg: #ffffff;
        --secondary-bg: #f8fafc;
        --sidebar-bg: #f1f5f9;
        --text-primary: #1e293b;
        --text-secondary: #475569;
        --border-color: #e2e8f0;
        --accent-color: #10b981;
        --hover-bg: #f1f5f9;
        --shadow: 0 1px 3px 0 rgba(0, 0, 0, 0.1), 0 1px 2px 0 rgba(0, 0, 0, 0.06);
    }
    
    /* Main App Styling */
    .stApp {
        background-color: var(--secondary-bg);
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
        line-height: 1.6;
    }
    
    /* Global text color fixes */
    .stApp, .stApp * {
        color: var(--text-primary) !important;
    }
    
    /* Main content area */
    .main .block-container {
        padding: 2rem 1rem;
        max-width: 1200px;
    }
    
    /* Titles and Headers */
    h1, h2, h3, h4, h5, h6 {
        color: var(--text-primary) !important;
        font-weight: 600 !important;
        margin-bottom: 1rem !important;
    }
    
    h1 {
        font-size: 2.25rem !important;
        border-bottom: 3px solid var(--accent-color);
        padding-bottom: 0.5rem;
    }
    
    /* Sidebar Styling */
    section[data-testid="stSidebar"] {
        background-color: var(--sidebar-bg) !important;
        border-right: 1px solid var(--border-color);
    }
    
    section[data-testid="stSidebar"] > div {
        padding: 1.5rem 1rem;
    }
    
    /* Sidebar text elements */
    section[data-testid="stSidebar"] h1,
    section[data-testid="stSidebar"] h2,
    section[data-testid="stSidebar"] h3,
    section[data-testid="stSidebar"] .stMarkdown,
    section[data-testid="stSidebar"] .stSelectbox label,
    section[data-testid="stSidebar"] p,
    section[data-testid="stSidebar"] span {
        color: var(--text-primary) !important;
        font-weight: 500;
    }
    
    /* Sidebar buttons */
    section[data-testid="stSidebar"] .stButton > button {
        background-color: var(--primary-bg) !important;
        color: var(--text-primary) !important;
        border: 1px solid var(--border-color) !important;
        border-radius: 0.5rem !important;
        padding: 0.5rem 1rem !important;
        font-weight: 500 !important;
        transition: all 0.2s ease !important;
        box-shadow: var(--shadow) !important;
        width: 100% !important;
        margin-bottom: 0.5rem !important;
    }
    
    section[data-testid="stSidebar"] .stButton > button:hover {
        background-color: var(--hover-bg) !important;
        border-color: var(--accent-color) !important;
        transform: translateY(-1px) !important;
    }
    
    /* Text inputs */
    .stTextInput > div > div > input {
        background-color: var(--primary-bg) !important;
        color: var(--text-primary) !important;
        border: 1px solid var(--border-color) !important;
        border-radius: 0.5rem !important;
        padding: 0.75rem !important;
        font-size: 1rem !important;
    }
    
    .stTextInput > div > div > input::placeholder {
        color: var(--text-secondary) !important;
    }
    
    /* Select boxes */
    .stSelectbox > div > div > select {
        background-color: var(--primary-bg) !important;
        color: var(--text-primary) !important;
        border: 1px solid var(--border-color) !important;
        border-radius: 0.5rem !important;
    }
    
    /* Primary buttons */
    .stButton > button[kind="primary"] {
        background-color: var(--accent-color) !important;
        color: white !important;
        border: none !important;
        border-radius: 0.5rem !important;
        padding: 0.75rem 2rem !important;
        font-weight: 600 !important;
        font-size: 1rem !important;
        transition: all 0.2s ease !important;
    }
    
    .stButton > button[kind="primary"]:hover {
        background-color: #059669 !important;
        transform: translateY(-1px) !important;
        box-shadow: 0 4px 12px rgba(16, 185, 129, 0.3) !important;
    }
    
    /* Chat Messages */
    .stChatMessage {
        background-color: var(--primary-bg) !important;
        border: 1px solid var(--border-color) !important;
        border-radius: 1rem !important;
        margin: 1rem 0 !important;
        box-shadow: var(--shadow) !important;
    }
    
    .stChatMessage[data-testid="chat-message-user"] {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%) !important;
        color: white !important;
        margin-left: 2rem !important;
    }
    
    .stChatMessage[data-testid="chat-message-user"] * {
        color: white !important;
    }
    
    .stChatMessage[data-testid="chat-message-assistant"] {
        background-color: var(--primary-bg) !important;
        margin-right: 2rem !important;
        border-left: 4px solid var(--accent-color) !important;
    }
    
    .stChatMessageContent {
        padding: 1rem 1.5rem !important;
        line-height: 1.6 !important;
        font-size: 1rem !important;
    }
    
    /* Chat input */
    .stChatInputContainer {
        background-color: var(--primary-bg) !important;
        border: 1px solid var(--border-color) !important;
        border-radius: 1rem !important;
        box-shadow: var(--shadow) !important;
        margin-top: 2rem !important;
    }
    
    .stChatInput > div > div > input {
        background-color: transparent !important;
        color: var(--text-primary) !important;
        border: none !important;
        font-size: 1rem !important;
        padding: 1rem !important;
    }
    
    .stChatInput > div > div > input::placeholder {
        color: var(--text-secondary) !important;
    }
    
    /* Alert messages */
    .stAlert {
        border-radius: 0.5rem !important;
        border: none !important;
        box-shadow: var(--shadow) !important;
    }
    
    .stAlert[data-testid="stAlert-info"] {
        background-color: #dbeafe !important;
        color: #1e40af !important;
        border-left: 4px solid #3b82f6 !important;
    }
    
    .stAlert[data-testid="stAlert-success"] {
        background-color: #dcfce7 !important;
        color: #166534 !important;
        border-left: 4px solid var(--accent-color) !important;
    }
    
    .stAlert[data-testid="stAlert-error"] {
        background-color: #fef2f2 !important;
        color: #dc2626 !important;
        border-left: 4px solid #ef4444 !important;
    }
    
    /* Spinner */
    .stSpinner {
        color: var(--accent-color) !important;
    }
    
    /* Dividers */
    hr {
        border: none !important;
        height: 1px !important;
        background-color: var(--border-color) !important;
        margin: 2rem 0 !important;
    }
    
    /* Links */
    a {
        color: var(--accent-color) !important;
        text-decoration: none !important;
        font-weight: 500 !important;
    }
    
    a:hover {
        text-decoration: underline !important;
        color: #059669 !important;
    }
    
    /* Citation links */
    .citation-link {
        background-color: #e0f2fe !important;
        color: #0277bd !important;
        padding: 0.25rem 0.5rem !important;
        border-radius: 0.25rem !important;
        font-size: 0.875rem !important;
        font-weight: 600 !important;
        text-decoration: none !important;
        display: inline-block !important;
        margin: 0.125rem !important;
        transition: all 0.2s ease !important;
    }
    
    .citation-link:hover {
        background-color: #b3e5fc !important;
        transform: translateY(-1px) !important;
    }
    
    /* Custom status box */
    .system-status {
        background: linear-gradient(135deg, #f0f9ff 0%, #e0f2fe 100%) !important;
        border: 1px solid #b3e5fc !important;
        color: var(--text-primary) !important;
        padding: 1.5rem !important;
        border-radius: 1rem !important;
        margin: 1.5rem 0 !important;
        box-shadow: var(--shadow) !important;
    }
    
    /* Scrollbar styling */
    ::-webkit-scrollbar {
        width: 6px;
        height: 6px;
    }
    
    ::-webkit-scrollbar-track {
        background: var(--secondary-bg);
    }
    
    ::-webkit-scrollbar-thumb {
        background: var(--border-color);
        border-radius: 3px;
    }
    
    ::-webkit-scrollbar-thumb:hover {
        background: #cbd5e1;
    }
    
    /* Responsive design */
    @media (max-width: 768px) {
        .main .block-container {
            padding: 1rem 0.5rem;
        }
        
        h1 {
            font-size: 1.75rem !important;
        }
        
        .stChatMessage[data-testid="chat-message-user"] {
            margin-left: 0.5rem !important;
        }
        
        .stChatMessage[data-testid="chat-message-assistant"] {
            margin-right: 0.5rem !important;
        }
    }
    
    /* Animation for typing effect */
    @keyframes blink {
        0%, 50% { opacity: 1; }
        51%, 100% { opacity: 0; }
    }
    
    .typing-cursor {
        animation: blink 1s infinite;
    }
    </style>
    """, unsafe_allow_html=True)
    
    try:
        enhanced_chatbot_interface()
    except Exception as e:
        st.error(f"Application error: {e}")
        st.info("Please refresh the page and try again.")
    
if __name__ == "__main__":
    main()
